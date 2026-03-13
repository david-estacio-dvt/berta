"""
knowledge_loader.py - Carga y extracción de texto de documentación CAS.

Lee todos los PDFs y DOCX de la carpeta documentation/ para construir
la base de conocimiento normativa que se inyecta en el prompt del agente.
"""

import os
import hashlib
import json
from pathlib import Path

# Directorio raíz del agente
AGENT_DIR = Path(__file__).parent
DOCS_DIR = AGENT_DIR / ".documentation"
CACHE_FILE = AGENT_DIR / ".knowledge_cache.json"


def _extract_pdf_text(filepath: str) -> str:
    """Extrae texto de un archivo PDF usando pypdf."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(filepath)
        texts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                texts.append(text.strip())
        return "\n\n".join(texts)
    except Exception as e:
        print(f"  [WARN] Error leyendo PDF {filepath}: {e}")
        return ""


def _extract_docx_text(filepath: str) -> str:
    """Extrae texto de un archivo DOCX usando python-docx."""
    try:
        from docx import Document
        doc = Document(filepath)
        texts = []
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text.strip())
        # También extraer tablas
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    texts.append(" | ".join(row_data))
        return "\n".join(texts)
    except Exception as e:
        print(f"  [WARN] Error leyendo DOCX {filepath}: {e}")
        return ""


def _file_hash(filepath: str) -> str:
    """Calcula hash MD5 del archivo para cache."""
    h = hashlib.md5()
    with open(filepath, 'rb') as f:
        for chunk in iter(lambda: f.read(8192), b''):
            h.update(chunk)
    return h.hexdigest()


def _load_cache() -> dict:
    """Carga la cache de conocimiento si existe."""
    if CACHE_FILE.exists():
        try:
            with open(CACHE_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            return {}
    return {}


def _save_cache(cache: dict):
    """Guarda la cache de conocimiento."""
    try:
        with open(CACHE_FILE, 'w', encoding='utf-8') as f:
            json.dump(cache, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"  [WARN] Error guardando cache: {e}")


def load_all_documents(max_chars_per_doc: int = 50000) -> str:
    """
    Carga y extrae texto de todos los documentos normativos en documentation/.
    
    Excluye la subcarpeta amv-documentacion-cas-prueba (casos de prueba reales).
    
    Args:
        max_chars_per_doc: Máximo de caracteres por documento individual.
        
    Returns:
        Texto concatenado de toda la documentación normativa.
    """
    if not DOCS_DIR.exists():
        print("[WARN] Carpeta .documentation/ no encontrada.")
        return "No se encontró documentación normativa."
    
    cache = _load_cache()
    updated = False
    
    all_texts = []
    
    # Archivos principales en documentation/ (excluir subcarpetas de prueba)
    files_to_process = []
    
    for item in sorted(DOCS_DIR.iterdir()):
        if item.is_file():
            ext = item.suffix.lower()
            if ext in ('.pdf', '.docx'):
                files_to_process.append(item)
    
    # También procesar la subcarpeta "Ejemplos ficheros"
    ejemplos_dir = DOCS_DIR / "Ejemplos ficheros"
    if ejemplos_dir.exists():
        for item in sorted(ejemplos_dir.iterdir()):
            if item.is_file():
                ext = item.suffix.lower()
                if ext in ('.xml', '.txt'):
                    files_to_process.append(item)
    
    print(f"\n[Knowledge Loader] Procesando {len(files_to_process)} documentos normativos...")
    
    for filepath in files_to_process:
        file_key = str(filepath)
        file_hash = _file_hash(str(filepath))
        
        # Comprobar cache
        if file_key in cache and cache[file_key].get('hash') == file_hash:
            text = cache[file_key]['text']
            print(f"  [CACHE] {filepath.name} ({len(text)} chars)")
        else:
            # Extraer texto
            ext = filepath.suffix.lower()
            if ext == '.pdf':
                text = _extract_pdf_text(str(filepath))
            elif ext == '.docx':
                text = _extract_docx_text(str(filepath))
            elif ext in ('.xml', '.txt'):
                try:
                    # Intentar diferentes encodings
                    for enc in ['utf-8', 'latin-1', 'utf-16']:
                        try:
                            with open(filepath, 'r', encoding=enc) as f:
                                text = f.read()
                            break
                        except (UnicodeDecodeError, UnicodeError):
                            continue
                    else:
                        text = ""
                except Exception as e:
                    print(f"  [WARN] Error leyendo {filepath.name}: {e}")
                    text = ""
            else:
                continue
            
            # Truncar si es muy largo
            if len(text) > max_chars_per_doc:
                text = text[:max_chars_per_doc] + "\n\n[... DOCUMENTO TRUNCADO ...]"
            
            # Actualizar cache
            cache[file_key] = {'hash': file_hash, 'text': text}
            updated = True
            print(f"  [NUEVO] {filepath.name} ({len(text)} chars)")
        
        if text:
            header = f"\n{'='*80}\n📄 DOCUMENTO: {filepath.name}\n{'='*80}\n"
            all_texts.append(header + text)
    
    if updated:
        _save_cache(cache)
    
    result = "\n\n".join(all_texts)
    print(f"[Knowledge Loader] Total: {len(result)} caracteres de conocimiento normativo.\n")
    
    return result


def load_test_case_files() -> list[dict]:
    """
    Lista los archivos de casos de prueba en amv-documentacion-cas-prueba/.
    
    Returns:
        Lista de dicts con {expediente, documento_id, extension, filepath}.
    """
    test_dir = DOCS_DIR / "amv-documentacion-cas-prueba"
    if not test_dir.exists():
        return []
    
    cases = []
    for item in sorted(test_dir.iterdir()):
        if item.is_file():
            # Formato de nombre: YYYYNNNNNNN-NNNNNNN-NNN.ext.ext
            name = item.stem
            parts = name.split('-')
            if len(parts) >= 2:
                cases.append({
                    'expediente': parts[0],
                    'documento_id': '-'.join(parts[:2]),
                    'extension': item.suffix.lower(),
                    'filepath': str(item),
                    'filename': item.name,
                    'size_kb': round(item.stat().st_size / 1024, 1)
                })
    
    return cases


if __name__ == "__main__":
    # Test standalone
    print("=== Cargando documentación normativa ===")
    kb = load_all_documents()
    print(f"\nResumen: {len(kb)} caracteres totales")
    print(f"Primeros 500 chars:\n{kb[:500]}")
    
    print("\n=== Casos de prueba disponibles ===")
    cases = load_test_case_files()
    for c in cases[:10]:
        print(f"  {c['expediente']} - {c['filename']} ({c['size_kb']} KB)")
    print(f"Total: {len(cases)} archivos de casos de prueba")
